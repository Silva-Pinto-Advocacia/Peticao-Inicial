"""
Silva Pinto Advocacia — Gerador Automático de Petições
Backend Flask com integração Claude API
Estratégia: extração de texto de todos os arquivos — zero PDFs nativos na API
"""

import os, sys, json, uuid, zipfile, shutil, logging, re, random, base64, gc
from pathlib import Path
from datetime import datetime

from flask import Flask, request, jsonify, send_file, render_template
import anthropic

# ── Config ──────────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024  # 200 MB

# ── Token budget ─────────────────────────────────────────────────────────────
# Sem limites artificiais — leitura completa de cada arquivo.
# Únicos limites são os físicos do modelo (200k tokens input ≈ 700k chars).
MAX_CHARS_PER_PDF      = 200_000   # essencialmente ilimitado para PDFs típicos
MAX_CHARS_PER_XLSX     = 200_000
MAX_CHARS_DOCX_MODEL   = 200_000   # modelo inteiro
MAX_CHARS_DOCX_OTHER   = 200_000   # relatórios e outros docx inteiros
MAX_PDFS               = 50        # praticamente todos os PDFs
TOTAL_CHAR_BUDGET      = 700_000   # ~175k tokens — margem segura abaixo do limite de 200k

# ── Helpers ──────────────────────────────────────────────────────────────────

def xe(text: str) -> str:
    return (text.replace("&","&amp;").replace("<","&lt;")
                .replace(">","&gt;").replace('"',"&quot;").replace("'","&apos;"))

def random_para_id() -> str:
    return f"{random.randint(0x10000000, 0x7FFFFFFE):08X}"

def extract_xml_text(xml: str) -> str:
    text = re.sub(r"<[^>]+>", " ", xml)
    return re.sub(r"\s+", " ", text).strip()

def truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n...[TRUNCADO — {len(text)-max_chars} chars omitidos]"

def extract_pdf_text(fpath: Path, max_chars: int = MAX_CHARS_PER_PDF) -> str:
    """Extract text from PDF using pypdf (lightweight, Python 3.14 safe)."""
    try:
        import pypdf
        text_parts = []
        with open(fpath, "rb") as f:
            reader = pypdf.PdfReader(f, strict=False)
            for page in reader.pages:
                try:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
                except Exception:
                    pass
        full = "\n".join(text_parts)
        if not full.strip():
            return f"[PDF sem texto extraível: {fpath.name}]"
        return truncate(full, max_chars)
    except Exception as e:
        log.warning("pypdf failed for %s: %s", fpath.name, e)
        return f"[Não foi possível extrair texto: {fpath.name}]"

def read_docx_text(path: Path, max_chars: int = MAX_CHARS_DOCX_OTHER) -> str:
    """Read text from docx using python-docx (no subprocess)."""
    try:
        import docx as _docx
        doc = _docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return truncate("\n".join(parts), max_chars)
    except Exception as e:
        log.warning("python-docx read error %s: %s", path.name, e)
        try:
            with zipfile.ZipFile(str(path)) as zf:
                xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
            return truncate(extract_xml_text(xml), max_chars)
        except Exception:
            return ""

def read_text_file(path: Path, max_chars: int = MAX_CHARS_PER_XLSX) -> str:
    try:
        # Try reading xlsx with openpyxl
        if path.suffix.lower() in (".xlsx", ".xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(path), data_only=True)
                parts = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    parts.append(f"[Planilha: {sheet}]")
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c) if c is not None else "" for c in row]
                        line = " | ".join(cells)
                        if line.strip(" |"):
                            parts.append(line)
                return truncate("\n".join(parts), max_chars)
            except Exception as e:
                log.warning("openpyxl failed: %s", e)
        with open(path, encoding="utf-8", errors="replace") as f:
            return truncate(f.read(), max_chars)
    except Exception:
        return ""


def parse_ficha_cliente(path: Path) -> dict:
    """Parse the structured 'Ficha do Cliente' XLSX to extract key fields.
    Returns a dict with computed fields like pontuacao_final = obtida + delta_anulacao.
    """
    out = {}

    # Phrases that indicate the field contains instructions, not real values
    PLACEHOLDER_PATTERNS = [
        r"^comarca\s+da\s+residencia",         # "comarca da residencia do cliente"
        r"^endere[çc]o\s+do\s+cliente",
        r"^(não|nao)\s+informad",              # "Não informado"
        r"^a\s+(definir|preencher|verificar)",
        r"^preencher",
        r"^marque\s+x",
        r"^se\s+sim",
        r"^sim\s*·\s*n[ãa]o",                  # "Sim · Não"
    ]

    def is_placeholder(value: str) -> bool:
        if not value:
            return True
        v = value.strip().lower()
        return any(re.match(p, v) for p in PLACEHOLDER_PATTERNS)

    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        ws = wb[wb.sheetnames[0]]

        kv = {}
        for row in ws.iter_rows(min_col=1, max_col=2, values_only=True):
            label = (row[0] or "").strip() if isinstance(row[0], str) else ""
            value = (row[1] or "").strip() if isinstance(row[1], str) else ""
            if not label and not value:
                continue
            if label:
                kv[label.lower()] = value

        def find_field(*keywords):
            """Find first label that contains all keywords. Skip placeholders."""
            for label, value in kv.items():
                if all(kw.lower() in label for kw in keywords) and value:
                    if is_placeholder(value):
                        continue
                    return value
            return ""

        # Extract structured fields
        out["nome_cliente"]      = find_field("nome", "cliente")
        out["concurso"]          = find_field("nome do concurso") or find_field("concurso *")
        out["banca"]             = find_field("banca")
        out["cargo"]             = find_field("cargo")
        out["comarca"]           = find_field("comarca")
        out["pontuacao_obtida"]  = find_field("pontuação obtida")
        out["nota_corte"]        = find_field("nota de corte")
        out["delta_anulacao"]    = find_field("pontuação possível após anulações") or find_field("após anulações")
        out["tipo_prova"]        = find_field("tipo da prova") or find_field("tipo de prova impugnada")
        out["questoes_anular"]   = find_field("questões a anular") or find_field("questoes a anular")
        out["gratuidade"]        = find_field("requer gratuidade")
        out["eliminado"]         = find_field("eliminado")
        out["resumo_fatos"]      = find_field("resumo dos fatos")
        out["observacoes"]       = find_field("observações") or find_field("instruções específicas")
        out["tipo_acao"]         = find_field("tipo de ação")
        out["proxima_fase"]      = find_field("data da próxima fase") or find_field("próxima fase")

        # Extract instituição/estado do concurso a partir do nome do concurso
        # Ex: "PCES – Polícia Civil do Espírito Santo" → instituição = Polícia Civil do Espírito Santo
        if out.get("concurso"):
            conc = out["concurso"]
            # Try to detect Polícia Civil + estado
            m = re.search(
                r"(Pol[íi]cia\s+(?:Civil|Militar|Federal|Rodovi[áa]ria)\s+(?:do\s+)?(?:Estado\s+(?:do|de|da)\s+)?[A-ZÁÉÍÓÚÂÊÔÃÕa-záéíóúâêôãõç ]+)",
                conc, re.IGNORECASE
            )
            if m:
                out["instituicao_concurso"] = m.group(1).strip().rstrip(",.")
            # Detect state code (ES, MG, RJ, SP etc.) at start of concurso
            m2 = re.match(r"^P[CMF][A-Z]{2}", conc.strip())
            if m2:
                state_code = m2.group(0)[2:]  # Last 2 letters
                state_map = {
                    "ES":"Espírito Santo", "MG":"Minas Gerais", "RJ":"Rio de Janeiro",
                    "SP":"São Paulo", "BA":"Bahia", "PR":"Paraná", "RS":"Rio Grande do Sul",
                    "PE":"Pernambuco", "CE":"Ceará", "GO":"Goiás", "DF":"Distrito Federal",
                    "AM":"Amazonas", "PA":"Pará", "MT":"Mato Grosso", "MS":"Mato Grosso do Sul",
                    "SC":"Santa Catarina", "AL":"Alagoas", "PB":"Paraíba", "RN":"Rio Grande do Norte",
                    "MA":"Maranhão", "PI":"Piauí", "TO":"Tocantins", "AC":"Acre",
                    "RO":"Rondônia", "RR":"Roraima", "AP":"Amapá", "SE":"Sergipe",
                }
                out["estado_concurso"] = state_map.get(state_code, state_code)
                out["uf_concurso"] = state_code
                # Sigla da instituição, ex: PCES
                out["sigla_concurso"] = m2.group(0)

        # Compute pontuacao_final: obtida + delta
        try:
            obtida_num = _extract_number(out["pontuacao_obtida"])
            delta_num  = _extract_number(out["delta_anulacao"])
            if obtida_num is not None and delta_num is not None:
                final = obtida_num + delta_num
                # Format as integer if whole, else with decimal
                final_str = str(int(final)) if final == int(final) else f"{final}"
                obtida_str = str(int(obtida_num)) if obtida_num == int(obtida_num) else f"{obtida_num}"
                delta_str  = str(int(delta_num))  if delta_num == int(delta_num)  else f"{delta_num}"
                out["pontuacao_final_calculada"] = f"{final_str} pontos"
                out["pontuacao_obtida_clean"]    = f"{obtida_str} pontos"
                out["delta_clean"]               = f"{delta_str} pontos"
                out["pontuacao_final_num"]       = final
                out["pontuacao_obtida_num"]      = obtida_num
                out["delta_num"]                 = delta_num
        except Exception:
            pass

        # Clean name: remove anything in parentheses
        if out.get("nome_cliente"):
            cleaned = re.sub(r"\s*\([^)]*\)", "", out["nome_cliente"]).strip()
            if cleaned:
                out["nome_cliente"] = cleaned

        # Extract structured personal data from "observações" (free text field)
        obs = out.get("observacoes", "")
        if obs:
            # RG: "RG MG16052778" or "RG nº 12345"
            m = re.search(r"RG\s*(?:n[ºo°.]*\s*)?([A-Z]{0,3}\s*[\d\.\-]+)", obs, re.I)
            if m: out["rg"] = m.group(1).strip()
            # CPF
            m = re.search(r"CPF\s*(?:n[ºo°.]*\s*)?(\d{3}\.\d{3}\.\d{3}-\d{2}|\d{11})", obs, re.I)
            if m: out["cpf"] = m.group(1).strip()
            # E-mail
            m = re.search(r"e-?mail\s*:?\s*([\w\.\-]+@[\w\.\-]+\.\w+)", obs, re.I)
            if m: out["email"] = m.group(1).strip()
            # Profissão
            m = re.search(r"profissão\s*:?\s*([^;,]+)", obs, re.I)
            if m: out["profissao"] = m.group(1).strip()
            # Estado civil
            m = re.search(r"estado civil\s*:?\s*([^;,]+)", obs, re.I)
            if m: out["estado_civil"] = m.group(1).strip()
            # Endereço — capture rua/número até CEP
            m = re.search(r"endere[çc]o\s*:?\s*(.+?)(?:CEP\s*([\d\-]+))?(?:\.|;|$)", obs, re.I)
            if m:
                addr = m.group(1).strip().rstrip(",")
                out["endereco_completo"] = addr
                if m.group(2):
                    out["cep"] = m.group(2).strip()
                # Try to extract city/UF from address: "..., Cidade/UF"
                cidade_match = re.search(r",\s*([\w\s]+?)\s*/\s*([A-Z]{2})", addr)
                if cidade_match:
                    out["cidade"] = cidade_match.group(1).strip()
                    out["uf"] = cidade_match.group(2).strip()
                    # Derive comarca from city (only if comarca itself is empty/placeholder)
                    if not out.get("comarca"):
                        out["comarca"] = f"{out['cidade'].upper()}/{out['uf']}"

        log.info("Ficha parseada: nome=%s, obtida=%s, delta=%s, final=%s, corte=%s",
                 out.get("nome_cliente"), out.get("pontuacao_obtida"),
                 out.get("delta_anulacao"), out.get("pontuacao_final_calculada"),
                 out.get("nota_corte"))
    except Exception as e:
        log.warning("Falha ao parsear ficha: %s", e)
    return out


def _extract_number(text: str) -> float | None:
    """Extract first numeric value from text like '66 pontos' or '6 pts'."""
    if not text:
        return None
    m = re.search(r"(\d+(?:[.,]\d+)?)", text)
    if m:
        try:
            return float(m.group(1).replace(",", "."))
        except ValueError:
            return None
    return None


_UF_TO_STATE = {
    "ES":"Espírito Santo", "MG":"Minas Gerais", "RJ":"Rio de Janeiro",
    "SP":"São Paulo", "BA":"Bahia", "PR":"Paraná", "RS":"Rio Grande do Sul",
    "PE":"Pernambuco", "CE":"Ceará", "GO":"Goiás", "DF":"Distrito Federal",
    "AM":"Amazonas", "PA":"Pará", "MT":"Mato Grosso", "MS":"Mato Grosso do Sul",
    "SC":"Santa Catarina", "AL":"Alagoas", "PB":"Paraíba", "RN":"Rio Grande do Norte",
    "MA":"Maranhão", "PI":"Piauí", "TO":"Tocantins", "AC":"Acre",
    "RO":"Rondônia", "RR":"Roraima", "AP":"Amapá", "SE":"Sergipe",
}


def _uf_to_state_name(uf: str) -> str:
    return _UF_TO_STATE.get((uf or "").upper(), uf or "")



def unpack_docx(docx_path: Path, out_dir: Path) -> bool:
    """Unpack docx (zip) into directory."""
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(str(docx_path), 'r') as zf:
            zf.extractall(str(out_dir))
        return True
    except Exception as e:
        log.error("unpack error: %s", e)
        return False

def pack_docx(unpacked_dir: Path, out_path: Path, original: Path) -> bool:
    """Repack directory into docx (zip), preserving original file list order."""
    try:
        # Get file list from original to preserve content types order
        orig_names = set()
        with zipfile.ZipFile(str(original), 'r') as orig_zf:
            orig_names = set(orig_zf.namelist())

        with zipfile.ZipFile(str(out_path), 'w', zipfile.ZIP_DEFLATED) as zf:
            # Write [Content_Types].xml first (required by OOXML spec)
            ct = unpacked_dir / "[Content_Types].xml"
            if ct.exists():
                zf.write(str(ct), "[Content_Types].xml")
            # Write all other files
            for fpath in sorted(unpacked_dir.rglob("*")):
                if fpath.is_file() and fpath.name != "[Content_Types].xml":
                    arcname = str(fpath.relative_to(unpacked_dir))
                    zf.write(str(fpath), arcname)
        return True
    except Exception as e:
        log.error("pack error: %s", e)
        return False

# ── System Prompt ─────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """Você é o Assistente Jurídico Especialista do escritório Silva Pinto Advocacia.
Advogado responsável: Dr. Casil da Silva Pinto — OAB/RJ nº 189.781.
Especialidade: Ações anulatórias de atos administrativos em concursos públicos.

SUA FUNÇÃO:
Analisar os arquivos do cliente, cruzar com o resumo do caso e com a ficha do cliente,
e gerar um JSON estruturado para preenchimento da Petição Inicial em modelo .docx.

═══ INSTRUÇÕES DE EXECUÇÃO ═══

1. ANÁLISE E EXTRAÇÃO DE DADOS
- Leia todos os arquivos extraídos do ZIP do cliente.
- Identifique os documentos de qualificação (procuração, RG, CPF, comprovante de residência)
  e extraia: nome completo, nacionalidade, estado civil, profissão, RG, CPF, e-mail e
  endereço completo com CEP.
- Para qualquer dado essencial não encontrado, use a marcação curta "[DADO AUSENTE]" e
  inclua no array "dados_ausentes" a descrição do que está faltando.

3. EXTRAÇÃO E SUBSTITUIÇÃO DAS QUESTÕES IMPUGNADAS
3.1) Identifique na "Ficha do Cliente" o campo "QUESTÕES A ANULAR" — essa lista
     define quais questões DEVEM aparecer na petição final do novo cliente.
3.2) No "RELATÓRIO TÉCNICO DAS QUESTÕES" (arquivo .docx separado), localize cada
     questão dessa lista. Para cada uma, extraia:
     - O número da questão
     - O vício (ERRO GROSSEIRO, EXTRAPOLAÇÃO DO EDITAL, DUPLICIDADE DE GABARITO etc.)
     - O enunciado e as alternativas
     - O resumo técnico-jurídico de uma frase
3.3) No MODELO da petição existe um capítulo que lista as questões antigas
     (do cliente anterior) — ex: "questões 10, 25, 31 e 34 da Prova Tipo 2".
     Você DEVE gerar pares de substituição para REMOVER os blocos das questões
     antigas e INSERIR os blocos das questões novas, mantendo a estrutura.
3.4) Para o capítulo "Da Probabilidade do Direito do Autor", o modelo traz o
     relatório técnico completo de UMA questão antiga. Você deve gerar um par
     de substituição que TROQUE esse relatório antigo pelo relatório íntegro
     de UMA das questões novas (escolha a primeira da lista, ou a com vício
     mais grave).
3.5) Popule o array "questoes" do JSON com cada questão nova:
     numero, vicio, resumo_peticao, enunciado, alternativas, gabarito_banca,
     resposta_correta, relatorio_integra (para a questão de destaque).
3.6) Em "questao_destaque_idx" indique o índice (0-based) da questão escolhida
     para o capítulo "Da Probabilidade do Direito".

REDAÇÃO JURÍDICA (Fatos, Direito, Pontuação)
- Linguagem altamente técnica, formal, persuasiva, adequada a Petição Inicial.
- Capítulo "Da Pontuação": MÁXIMO 3 parágrafos, indicando a pontuação que o
  candidato alcançará após anulação e dando ÊNFASE em que atingirá a nota de corte.
- Use SEMPRE a "PONTUAÇÃO FINAL APÓS ANULAÇÕES" calculada nos DADOS AUTORITATIVOS.
  Nunca invente outro número.
- Se o cliente NÃO tiver direito à gratuidade conforme a ficha, marque
  "gratuidade": false — o sistema removerá o capítulo correspondente.

4. ROL DE DOCUMENTOS
- Liste em "rol_documentos" todos os documentos que devem instruir a petição,
  numerados sequencialmente. Para cada item, indique o nome do arquivo correspondente
  no ZIP (campo "arquivo_correspondente") para que o sistema renomeie automaticamente.
- Inclua RG, CPF, Procuração, Comprovante de Residência, Cartão Resposta,
  Edital, e os Pareceres/Relatórios Técnicos de cada questão a ser anulada.

5. CONFERÊNCIAS OBRIGATÓRIAS
- O número de questões em "questoes" deve coincidir exatamente com as questões
  pedidas para anulação na ficha do candidato.
- A comarca em "processo.comarca" deve ser a do domicílio do cliente.
- Use o gênero correto do cliente em toda a redação (concordância nominal e verbal).
- Não invente jurisprudências — só use as que constam no modelo enviado.

═══ REGRAS INVIOLÁVEIS ═══

- Use somente dados reais extraídos dos documentos anexados — nunca invente.
- Não use jurisprudências ou citações genéricas; use apenas as do modelo.
- Diferencie petições por procedimento:
  * Procedimento Comum: o rol completo de questões já vai na inicial.
  * Tutela Cautelar Antecedente: pode informar que o rol completo virá em emenda.
- Todas as informações do candidato devem refletir os documentos do cliente atual,
  em substituição completa aos dados que constam no modelo.

═══ SUBSTITUIÇÃO DOS DADOS DO CLIENTE NO MODELO ═══

ATENÇÃO MÁXIMA: Esta é a tarefa MAIS CRÍTICA da sua função.
O modelo da petição vem com dados de um CLIENTE ANTIGO escritos diretamente no texto.
Você DEVE identificar TODOS esses trechos sem exceção e gerar pares
"buscar → substituir" para que o sistema faça a troca automaticamente.

TIPOS DE DADOS QUE PRECISAM SER SUBSTITUÍDOS (lista NÃO exaustiva — identifique tudo):

1. QUALIFICAÇÃO PESSOAL DO CLIENTE
   - Nome completo (em TODAS as ocorrências, inclusive em letras maiúsculas)
   - Estado civil (ex: "Casado", "Solteira")
   - Profissão (ex: "desempregado", "policial militar")
   - Número do RG (ex: "339924-4")
   - Número do CPF (ex: "140.948.387-83")
   - E-mail
   - Endereço completo (rua, número, bairro, complemento)
   - Cidade e UF
   - CEP

2. DADOS DO CONCURSO
   - Nome do concurso e edital (ex: "Edital nº 01/2025 — PCES")
   - Cargo pretendido (ex: "Oficial Investigador de Polícia")
   - Banca examinadora (ex: "IBADE", "FGV", "AOCP")
   - Tipo de prova realizada (ex: "Prova Tipo 2")
   - Datas das etapas

3. PONTUAÇÃO E CLASSIFICAÇÃO
   - Pontuação obtida pelo cliente antigo (ex: "55 pontos", "55 pts")
   - Pontuação após anulação (ex: "59 pts", "59 pontos")
   - Nota de corte (ex: "58 pontos") — se diferente para o novo cliente
   - Quaisquer outros números relacionados à pontuação

4. JURISDIÇÃO E PARTES
   - Comarca de endereçamento (ex: "BARRA DE SÃO FRANCISCO/ES")
   - Estado/UF que compõe o polo passivo
   - Tipo de ação (se diferente)

5. QUESTÕES IMPUGNADAS
   - Lista de números das questões antigas (ex: "questões 10, 25, 31 e 34")
   - Substituir pela lista das questões reais do novo cliente

REGRAS OBRIGATÓRIAS PARA OS PARES:

✅ PARES GRANULARES: Crie um par por dado individual em vez de um par gigante.
   - BOM: {"buscar": "ADLER MARQUES DE LIMA", "substituir": "JOÃO DA SILVA"}
   - BOM: {"buscar": "140.948.387-83", "substituir": "111.222.333-44"}
   - BOM: {"buscar": "55 pontos", "substituir": "60 pontos"}
   - RUIM: {"buscar": "ADLER MARQUES DE LIMA, Casado, desempregado, RG 339924-4...",
            "substituir": "JOÃO DA SILVA, Solteiro, professor, RG 555..."}

✅ TEXTO LITERAL: copie EXATAMENTE como aparece no modelo (acentos, maiúsculas,
   pontuação, espaços). Se aparece "ADLER MARQUES DE LIMA" em maiúsculas,
   crie um par com texto em maiúsculas. Se aparece "Adler Marques de Lima"
   em outro lugar, crie OUTRO par para essa variação.

✅ BUSQUE TODAS AS OCORRÊNCIAS: percorra o modelo inteiro mentalmente e
   procure cada dado em CADA seção (cabeçalho, qualificação, fatos, pontuação,
   pedidos, etc.). NUNCA assuma que substituir uma vez resolve tudo —
   se o sistema vai trocar todas as ocorrências de uma string, mas se o nome
   aparece em formatos diferentes (maiúsculo, minúsculo, abreviado), cada
   formato precisa de seu próprio par.

✅ NÃO INCLUA DUPLICATAS: se a string idêntica aparece 5 vezes, basta UM par.

✅ MÍNIMO 15 PARES: para um modelo típico de petição de concurso, você deve
   gerar AO MENOS 15-25 pares de substituição. Se gerar menos de 10, é sinal
   de que está deixando dados antigos passar.

JSON PURO (sem markdown, sem backticks, sem texto antes ou depois).
REGRAS CRÍTICAS PARA O JSON:
- Use APENAS aspas duplas para strings, nunca aspas simples ou aspas tipográficas.
- Dentro de strings, escape quebras de linha como \\n (nunca quebra de linha real).
- Dentro de strings, escape aspas duplas como \\".
- NÃO inclua vírgulas após o último item de objetos ou arrays.
- Não use comentários no JSON.
- Mantenha valores curtos e objetivos para evitar exceder o limite de tokens.
- Antes de responder, valide mentalmente a sintaxe do JSON.

Schema:
{
  "cliente": {
    "nome_completo": "",
    "nacionalidade": "",
    "estado_civil": "",
    "profissao": "",
    "rg": "",
    "cpf": "",
    "email": "",
    "endereco": "",
    "cidade": "",
    "uf": "",
    "cep": "",
    "genero": "M ou F"
  },
  "processo": {
    "tipo_acao": "",
    "comarca": "",
    "banca": "",
    "cargo": "",
    "concurso": "",
    "pontuacao_obtida": "",
    "pontuacao_corte": "",
    "pontuacao_apos_anulacao": "",
    "categoria": "ampla_concorrencia ou cotas",
    "gratuidade": true,
    "procedimento": "comum ou cautelar_antecedente"
  },
  "questoes": [
    {
      "numero": 0,
      "vicio": "",
      "resumo_peticao": "",
      "enunciado": "",
      "alternativas": ["A) ...", "B) ...", "C) ...", "D) ...", "E) ..."],
      "gabarito_banca": "",
      "resposta_correta": "",
      "relatorio_integra": ""
    }
  ],
  "questao_destaque_idx": 0,
  "substituicoes": [
    {"buscar": "ADLER MARQUES DE LIMA", "substituir": "NOME DO NOVO CLIENTE"},
    {"buscar": "55 pontos", "substituir": "60 pontos"}
  ],
  "textos": {
    "fatos": "",
    "pontuacao": "",
    "probabilidade_direito": "",
    "fundamentos_juridicos": ""
  },
  "rol_documentos": [
    {"numero": 1, "descricao": "", "arquivo_correspondente": ""}
  ],
  "dados_ausentes": [],
  "relatorio_alteracoes": []
}"""

# ── Claude call ───────────────────────────────────────────────────────────────

def _parse_claude_json(raw: str) -> dict:
    """Robustly parse JSON from Claude's response, handling common malformations."""
    # Strip markdown fences
    raw = re.sub(r"```json\s*", "", raw)
    raw = re.sub(r"```\s*", "", raw)
    raw = raw.strip()

    # Extract just the JSON object — find the outermost { ... }
    start = raw.find("{")
    end   = raw.rfind("}")
    if start >= 0 and end > start:
        raw = raw[start:end+1]

    # First attempt: parse as-is
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e1:
        log.warning("JSON decode failed at first try: %s", e1)

    # Second attempt: fix common issues
    fixed = raw
    # Remove trailing commas before } or ]
    fixed = re.sub(r",(\s*[}\]])", r"\1", fixed)
    # Replace smart quotes that may have leaked in
    fixed = fixed.replace("\u201c", '"').replace("\u201d", '"')
    fixed = fixed.replace("\u2018", "'").replace("\u2019", "'")
    # Newlines inside strings — most common cause of "Expecting ',' delimiter"
    # is unescaped newlines in JSON string values. Try escaping them.
    try:
        return json.loads(fixed)
    except json.JSONDecodeError as e2:
        log.warning("JSON decode failed at second try: %s", e2)

    # Third attempt: aggressive — try to fix unescaped newlines inside strings
    # by parsing line-by-line and joining quoted strings
    try:
        # Replace literal newlines inside quoted strings with \\n
        result = []
        in_string = False
        escape = False
        for ch in fixed:
            if escape:
                result.append(ch)
                escape = False
                continue
            if ch == "\\":
                result.append(ch)
                escape = True
                continue
            if ch == '"':
                in_string = not in_string
                result.append(ch)
                continue
            if in_string and ch == "\n":
                result.append("\\n")
                continue
            if in_string and ch == "\r":
                result.append("\\r")
                continue
            if in_string and ch == "\t":
                result.append("\\t")
                continue
            result.append(ch)
        fixed = "".join(result)
        return json.loads(fixed)
    except json.JSONDecodeError as e3:
        log.error("All JSON parse attempts failed. Last error: %s", e3)
        log.error("Raw content (first 500 chars): %s", raw[:500])
        log.error("Raw content (last 500 chars): %s", raw[-500:])
        raise


def call_claude(api_key: str, full_text: str) -> dict:
    """Single text-only call to Claude. No PDFs, no binary — pure text."""
    # 5-minute timeout: longer than gunicorn's, so Claude has time to respond
    client = anthropic.Anthropic(api_key=api_key, timeout=300.0, max_retries=2)
    log.info("Sending %d chars to Claude", len(full_text))
    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=64000,  # max para Claude Sonnet 4.5
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": full_text}]
    )
    log.info("Claude response: %d input / %d output tokens. Stop reason: %s",
             message.usage.input_tokens if message.usage else 0,
             message.usage.output_tokens if message.usage else 0,
             message.stop_reason)

    # Detect truncation
    if message.stop_reason == "max_tokens":
        raise ValueError(
            "A resposta do Claude foi truncada por exceder o limite de tokens. "
            "Tente um ZIP com menos arquivos ou peça em duas etapas."
        )

    raw = "".join(b.text for b in message.content if hasattr(b, "text"))
    return _parse_claude_json(raw)

# ── DOCX editing ──────────────────────────────────────────────────────────────

def apply_substitutions(unpacked_dir: Path, data: dict) -> list[str]:
    """Apply find/replace pairs from Claude across XML and headers/footers."""
    changes = []

    # Files to edit: document.xml + headers + footers
    word_dir = unpacked_dir / "word"
    files_to_edit = []
    if (word_dir / "document.xml").exists():
        files_to_edit.append(word_dir / "document.xml")
    for f in word_dir.glob("header*.xml"):
        files_to_edit.append(f)
    for f in word_dir.glob("footer*.xml"):
        files_to_edit.append(f)

    if not files_to_edit:
        return ["ERRO: nenhum XML editável encontrado"]

    # ── STEP A: Rewrite "ROL DE QUESTÕES ILEGAIS" chapter FIRST ──────────────
    # Important to do this BEFORE pair substitutions, so that pairs targeting
    # old question text don't accidentally match new content (and vice-versa).
    questoes = data.get("questoes", [])
    if questoes:
        doc_xml_path = unpacked_dir / "word" / "document.xml"
        if doc_xml_path.exists():
            xml = doc_xml_path.read_text(encoding="utf-8")
            new_xml, n_replaced = _rewrite_questoes_chapter(xml, questoes)
            if n_replaced > 0:
                doc_xml_path.write_text(new_xml, encoding="utf-8")
                changes.append(
                    f"📜 Capítulo 'Rol de Questões Ilegais' reescrito com {len(questoes)} questões novas"
                )
            else:
                changes.append(
                    "ℹ️ Capítulo 'Rol de Questões Ilegais' não localizado para reescrita automática"
                )

    # ── STEP B: Apply find/replace pairs from Claude response ────────────────
    pairs = data.get("substituicoes", [])
    log.info("Total de substituições a aplicar: %d", len(pairs))
    if not pairs:
        changes.append("⚠️ Nenhum par de substituição retornado pelo Claude")

    # Apply each pair across all XML files
    for pair_idx, pair in enumerate(pairs):
        old = pair.get("buscar", "").strip()
        new = pair.get("substituir", "").strip()
        if not old:
            log.warning("Par #%d: 'buscar' vazio, pulando", pair_idx)
            continue
        if old == new:
            log.info("Par #%d: buscar==substituir, pulando", pair_idx)
            continue
        # CRITICAL SAFETY: if 'old' is contained in 'new', applying repeatedly
        # would cause an infinite loop appending the same content forever.
        if old in new:
            log.warning(
                "Par #%d: 'buscar' está contido em 'substituir' — par recursivo, "
                "pulando para evitar loop. buscar='%s', substituir='%s'",
                pair_idx, old[:60], new[:60]
            )
            changes.append(
                f"⚠️ Par recursivo ignorado (geraria loop): '{old[:60]}' → '{new[:60]}'"
            )
            continue

        log.info("Par #%d: buscando '%s' (len=%d)", pair_idx, old[:80], len(old))

        old_xe = xe(old)
        new_xe = xe(new)
        total_count = 0
        method_used = ""

        for xml_path in files_to_edit:
            xml = xml_path.read_text(encoding="utf-8")
            file_count_literal = 0
            file_count_crossrun = 0

            # Method 1: literal XML-escaped match (substitutes ALL occurrences)
            count = xml.count(old_xe)
            if count > 0:
                xml = xml.replace(old_xe, new_xe)
                file_count_literal = count

            # Method 2: cross-run replacement (loop for any remaining fragmented occurrences)
            attempts = 0
            while attempts < 30:  # safety limit
                new_xml = _replace_across_runs(xml, old, new)
                if new_xml == xml:
                    break
                xml = new_xml
                file_count_crossrun += 1
                attempts += 1

            file_count = file_count_literal + file_count_crossrun
            if file_count > 0:
                xml_path.write_text(xml, encoding="utf-8")
                method_used = "mixed"
                if file_count_literal > 0 and file_count_crossrun == 0:
                    method_used = "literal"
                elif file_count_crossrun > 0 and file_count_literal == 0:
                    method_used = "cross-run"
                log.info("  %s: %d replacement(s) via %s (lit=%d, x-run=%d)",
                         xml_path.name, file_count, method_used,
                         file_count_literal, file_count_crossrun)

            total_count += file_count

        if total_count > 0:
            changes.append(f"✅ '{old[:50]}' → '{new[:50]}' ({total_count}x)")
        else:
            # "Não encontrado" não é necessariamente um erro — pode ser que o
            # texto sequer existia no modelo, o que é normal.
            log.info("  Texto não encontrado (pode estar correto): '%s'", old[:80])
            changes.append(f"ℹ️ Não localizado (talvez já correto): '{old[:60]}'")

    # Remove gratuidade chapter if not applicable
    p = data.get("processo", {})
    if not p.get("gratuidade", True):
        for xml_path in files_to_edit:
            xml = xml_path.read_text(encoding="utf-8")
            new_xml, removed = _remove_gratuidade_chapter(xml)
            if removed:
                xml_path.write_text(new_xml, encoding="utf-8")
                changes.append("🗑️ Capítulo de gratuidade removido")
                break

    return changes


def _rewrite_questoes_chapter(xml: str, questoes: list) -> tuple[str, int]:
    """Locate the chapter 'DO ROL DE QUESTÕES ILEGAIS' (or similar) and replace
    its question blocks with new ones based on the questoes list from IA.

    Strategy: find the chapter title paragraph, then walk forward until the next
    chapter title (typically all-caps line) and replace everything between with
    new question blocks.
    """
    if not questoes:
        return xml, 0

    # Patterns that mark the start of the chapter
    chapter_patterns = [
        r"DO ROL DE QUEST[ÕO]ES ILEGAIS",
        r"ROL DE QUEST[ÕO]ES ILEGAIS",
        r"DAS QUEST[ÕO]ES ILEGAIS",
        r"DA ILEGALIDADE DAS QUEST[ÕO]ES",
        r"DAS QUEST[ÕO]ES ANUL[ÁA]VEIS",
    ]
    # Find chapter start in the visible text
    pattern = re.compile(r'<w:t[^>]*>([^<]*)</w:t>', re.DOTALL)
    matches = list(pattern.finditer(xml))
    if not matches:
        return xml, 0

    concat = ""
    positions = []  # (start_in_concat, end_in_concat, match_idx)
    for i, m in enumerate(matches):
        text = m.group(1)
        start = len(concat)
        concat += text
        positions.append((start, start + len(text), i))

    chapter_start_idx = None
    for cp in chapter_patterns:
        m = re.search(cp, concat, re.IGNORECASE)
        if m:
            chapter_start_idx = m.start()
            break
    if chapter_start_idx is None:
        return xml, 0

    # Find the paragraph containing chapter_start_idx
    chapter_start_match_idx = None
    for s, e, i in positions:
        if s <= chapter_start_idx < e:
            chapter_start_match_idx = i
            break
    if chapter_start_match_idx is None:
        return xml, 0

    # Find the corresponding <w:p> element start of that paragraph
    chapter_start_in_xml = matches[chapter_start_match_idx].start()
    p_start = xml.rfind('<w:p ', 0, chapter_start_in_xml)
    if p_start < 0:
        p_start = xml.rfind('<w:p>', 0, chapter_start_in_xml)
    if p_start < 0:
        return xml, 0
    # End of chapter heading paragraph
    p_end_marker = xml.find('</w:p>', p_start) + len('</w:p>')

    # Now find the END of the chapter — next paragraph that looks like another chapter title
    # (heuristic: ALL CAPS title >= 8 chars, mostly uppercase letters)
    # Walk through paragraphs after chapter heading
    end_of_chapter_in_xml = p_end_marker  # default: just after heading
    next_chapter_keywords = [
        r"DO\s+M[ÉE]RITO", r"DA\s+TUTELA", r"DOS\s+REQUERIMENTOS",
        r"DOS\s+PEDIDOS", r"DA\s+ASSIST[ÊE]NCIA", r"DA\s+GRATUIDADE",
        r"DO\s+VALOR\s+DA\s+CAUSA", r"DAS\s+PROVAS", r"DO\s+R[ÉE]U",
        r"DO\s+JUIZADO", r"REQUER", r"TERMOS\s+EM\s+QUE",
        r"DA\s+CITA[ÇC][ÃA]O", r"DA\s+ANULA[ÇC][ÃA]O\s+DAS\s+QUEST",  # explicit chapter
    ]
    next_chapter_re = re.compile("|".join(next_chapter_keywords), re.IGNORECASE)

    # Find next chapter title in concat after chapter_start_idx + chapter_title_len
    # We skip 100 chars to avoid the chapter title itself
    search_start = chapter_start_idx + 50
    next_match = next_chapter_re.search(concat, search_start)
    if next_match:
        next_match_idx_in_concat = next_match.start()
        # Find which match index corresponds
        for s, e, i in positions:
            if s <= next_match_idx_in_concat < e:
                next_p_start_in_xml = xml.rfind('<w:p ', 0, matches[i].start())
                if next_p_start_in_xml < 0:
                    next_p_start_in_xml = xml.rfind('<w:p>', 0, matches[i].start())
                if next_p_start_in_xml > p_end_marker:
                    end_of_chapter_in_xml = next_p_start_in_xml
                break

    # Build new XML for the chapter content (after the heading paragraph)
    new_blocks_xml = ""
    for q in questoes:
        numero = q.get("numero", "?")
        vicio  = q.get("vicio", "VÍCIO NÃO ESPECIFICADO")
        enunciado = q.get("enunciado", "")
        alternativas = q.get("alternativas", []) or []
        relatorio = q.get("relatorio_integra") or q.get("resumo_peticao", "")

        # Header: QUESTÃO N — VÍCIO
        new_blocks_xml += (
            f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:pPr><w:spacing w:before="240" w:after="120"/></w:pPr>'
            f'<w:r><w:rPr><w:b/><w:caps/></w:rPr>'
            f'<w:t xml:space="preserve">QUESTÃO {xe(str(numero))} — {xe(str(vicio))}</w:t>'
            f'</w:r></w:p>'
        )

        # Enunciado (single paragraph with command + statement)
        if enunciado:
            new_blocks_xml += (
                f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f'<w:pPr><w:jc w:val="both"/></w:pPr>'
                f'<w:r><w:t xml:space="preserve">{xe(enunciado)}</w:t></w:r>'
                f'</w:p>'
            )

        # Alternativas (one per line)
        for alt in alternativas:
            if alt:
                new_blocks_xml += (
                    f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    f'<w:pPr><w:ind w:left="567"/></w:pPr>'
                    f'<w:r><w:t xml:space="preserve">{xe(str(alt))}</w:t></w:r>'
                    f'</w:p>'
                )

        # Fundamentação completa (preferable over short summary)
        if relatorio:
            # Split by paragraphs if it has line breaks
            for paragraph in str(relatorio).split("\n"):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                new_blocks_xml += (
                    f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    f'<w:pPr><w:jc w:val="both"/></w:pPr>'
                    f'<w:r><w:t xml:space="preserve">{xe(paragraph)}</w:t></w:r>'
                    f'</w:p>'
                )

    # Replace chapter content (between heading end and next chapter start)
    new_xml = xml[:p_end_marker] + new_blocks_xml + xml[end_of_chapter_in_xml:]
    return new_xml, len(questoes)



def _replace_across_runs(xml: str, old: str, new: str) -> str:
    """Try to replace text that may be split across multiple <w:t> elements."""
    if not old:
        return xml

    # Extract all visible text spans
    pattern = re.compile(r'<w:t[^>]*>([^<]*)</w:t>', re.DOTALL)
    matches = list(pattern.finditer(xml))
    if not matches:
        return xml

    # Build concatenated text and position map
    concat = ""
    positions = []
    for i, m in enumerate(matches):
        text = m.group(1)
        start = len(concat)
        concat += text
        positions.append((start, start + len(text), i))

    # Try direct find first
    idx = concat.find(old)

    # Whitespace-tolerant fallback: collapse whitespace in both
    if idx < 0:
        norm_old = re.sub(r"\s+", " ", old).strip()
        norm_concat = re.sub(r"\s+", " ", concat).strip()
        if norm_old in norm_concat:
            # Find the actual range in concat by walking char by char
            # building a normalized version while tracking original positions
            def find_norm(haystack: str, needle: str):
                """Returns (start_in_haystack, end_in_haystack) for a whitespace-normalised match."""
                # Build mapping of original_index -> normalized_index
                norm_chars = []
                orig_indices = []
                last_was_space = True
                for i, ch in enumerate(haystack):
                    if ch.isspace():
                        if not last_was_space:
                            norm_chars.append(" ")
                            orig_indices.append(i)
                            last_was_space = True
                    else:
                        norm_chars.append(ch)
                        orig_indices.append(i)
                        last_was_space = False
                norm_str = "".join(norm_chars).strip()
                # Find needle (also normalised)
                n_idx = norm_str.find(needle)
                if n_idx < 0:
                    return None
                # Account for leading whitespace stripped
                leading_strip = len(norm_str) - len(norm_str.lstrip()) if False else 0
                # Map back to original
                start_orig = orig_indices[n_idx] if n_idx < len(orig_indices) else None
                end_norm   = n_idx + len(needle) - 1
                end_orig   = orig_indices[end_norm] + 1 if end_norm < len(orig_indices) else len(haystack)
                return (start_orig, end_orig)

            result = find_norm(concat, norm_old)
            if result:
                idx, end_idx_calc = result
            else:
                return xml
        else:
            return xml
    else:
        end_idx_calc = idx + len(old)

    end_idx = end_idx_calc

    # Identify runs that overlap [idx, end_idx)
    affected = [(s, e, i) for (s, e, i) in positions if not (e <= idx or s >= end_idx)]
    if not affected:
        return xml

    affected_idxs = [a[2] for a in affected]
    first_affected = affected[0][2]

    # Build new XML
    new_xml_parts = []
    last_pos = 0
    for j, m in enumerate(matches):
        new_xml_parts.append(xml[last_pos:m.start()])
        if j == first_affected:
            run_start, run_end, _ = positions[j]
            prefix = m.group(1)[:max(0, idx - run_start)]
            suffix = m.group(1)[max(0, end_idx - run_start):] if run_end > end_idx else ""
            new_text = prefix + new + suffix
            new_xml_parts.append(f'<w:t xml:space="preserve">{xe(new_text)}</w:t>')
        elif j in affected_idxs:
            run_start, run_end, _ = positions[j]
            if run_end <= end_idx:
                new_xml_parts.append('<w:t xml:space="preserve"></w:t>')
            else:
                suffix = m.group(1)[end_idx - run_start:]
                new_xml_parts.append(f'<w:t xml:space="preserve">{xe(suffix)}</w:t>')
        else:
            new_xml_parts.append(m.group(0))
        last_pos = m.end()
    new_xml_parts.append(xml[last_pos:])
    return "".join(new_xml_parts)


def _build_qualificacao(c: dict) -> str:
    parts = []
    if c.get("nacionalidade"): parts.append(c["nacionalidade"])
    if c.get("estado_civil"):  parts.append(c["estado_civil"])
    if c.get("profissao"):     parts.append(c["profissao"])
    if c.get("rg"):            parts.append(f"portador do RG nº {c['rg']}")
    if c.get("cpf"):           parts.append(f"inscrito no CPF/MF sob o nº {c['cpf']}")
    if c.get("email"):         parts.append(f"e-mail {c['email']}")
    addr = _build_endereco(c)
    if addr:                   parts.append(f"residente e domiciliado em {addr}")
    return ", ".join(parts)

def _build_endereco(c: dict) -> str:
    return ", ".join(p for p in [c.get("endereco",""), c.get("cidade",""), c.get("uf",""), c.get("cep","")] if p)

def _replace_block(xml: str, marker: str, new_text: str, changes: list) -> str:
    if marker not in xml or not new_text:
        return xml
    xml = xml.replace(marker, _text_to_paragraphs(new_text))
    changes.append(f"✅ Bloco: {marker}")
    return xml

def _text_to_paragraphs(text: str) -> str:
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    result = []
    for para in paras:
        pid = random_para_id()
        result.append(
            f'<w:p w14:paraId="{pid}" w14:textId="FFFFFFFF" w:rsidR="00000000">'
            f'<w:pPr><w:ind w:left="0"/><w:jc w:val="both"/></w:pPr>'
            f'<w:r><w:t xml:space="preserve">{xe(para.replace(chr(10)," "))}</w:t></w:r>'
            f'</w:p>'
        )
    return "\n".join(result)

def _insert_questoes(xml: str, questoes: list, changes: list) -> str:
    if not questoes:
        return xml
    marker = next((m for m in ["BLOCO_QUESTOES_ILEGAIS","DAS_QUESTOES_ILEGAIS","QUESTOES_ANULAVEIS"] if m in xml), None)
    if not marker:
        changes.append("⚠️ Marcador de questões não encontrado no modelo")
        return xml

    blocks = []
    for q in sorted(questoes, key=lambda x: x.get("numero", 0)):
        num, vicio = q.get("numero","?"), q.get("vicio","VÍCIO").upper()
        pid = random_para_id()
        blocks.append(
            f'<w:p w14:paraId="{pid}" w14:textId="FFFFFFFF" w:rsidR="00000000">'
            f'<w:pPr><w:jc w:val="both"/><w:rPr><w:b/></w:rPr></w:pPr>'
            f'<w:r><w:rPr><w:b/></w:rPr><w:t>{xe(f"QUESTÃO {num} — {vicio}")}</w:t></w:r></w:p>'
        )
        if q.get("enunciado"):
            pid = random_para_id()
            blocks.append(
                f'<w:p w14:paraId="{pid}" w14:textId="FFFFFFFF" w:rsidR="00000000">'
                f'<w:pPr><w:jc w:val="both"/></w:pPr>'
                f'<w:r><w:t xml:space="preserve">{xe(q["enunciado"])}</w:t></w:r></w:p>'
            )
        for alt in q.get("alternativas", []):
            pid = random_para_id()
            blocks.append(
                f'<w:p w14:paraId="{pid}" w14:textId="FFFFFFFF" w:rsidR="00000000">'
                f'<w:pPr><w:ind w:left="720"/><w:jc w:val="both"/></w:pPr>'
                f'<w:r><w:t xml:space="preserve">{xe(alt)}</w:t></w:r></w:p>'
            )
        if q.get("resumo_peticao"):
            pid = random_para_id()
            blocks.append(
                f'<w:p w14:paraId="{pid}" w14:textId="FFFFFFFF" w:rsidR="00000000">'
                f'<w:pPr><w:jc w:val="both"/></w:pPr>'
                f'<w:r><w:t xml:space="preserve">{xe(q["resumo_peticao"])}</w:t></w:r></w:p>'
            )
        pid = random_para_id()
        blocks.append(f'<w:p w14:paraId="{pid}" w14:textId="FFFFFFFF" w:rsidR="00000000"><w:pPr/></w:p>')

    xml = xml.replace(marker, "\n".join(blocks))
    changes.append(f"✅ {len(questoes)} questão(ões) inserida(s)")
    return xml

def _remove_gratuidade_chapter(xml: str):
    pat = r'<w:p[^>]*>.*?gratuidade.*?</w:p>\s*(<w:p[^>]*>.*?</w:p>\s*)*?(?=<w:p[^>]*>[^<]*(?:DA PROBABILIDADE|DO DIREITO|DOS PEDIDOS))'
    new_xml, n = re.subn(pat, "", xml, flags=re.IGNORECASE | re.DOTALL)
    return (new_xml, True) if n else (xml, False)

# ── File renaming ─────────────────────────────────────────────────────────────

def rename_files_by_rol(work_dir: Path, rol: list) -> dict:
    mapping = {}
    files_in_dir = {f.name.lower(): f for f in work_dir.iterdir() if f.is_file()}
    for item in rol:
        num  = item.get("numero", "")
        desc = item.get("descricao", "")
        orig = item.get("arquivo_correspondente", "")
        safe = re.sub(r"\s+", "_", re.sub(r"[^\w\s\-]", "", desc).strip())
        src  = None
        if orig:
            candidate = work_dir / orig
            if candidate.exists():
                src = candidate
        if not src:
            orig_l = orig.lower()
            for fl, fp in files_in_dir.items():
                if orig_l in fl or fl in orig_l:
                    src = fp
                    break
        if src and src.exists():
            new_name = f"{num}. {safe}{src.suffix}"
            src.rename(work_dir / new_name)
            mapping[src.name] = new_name
        else:
            mapping[f"[AUSENTE] {orig}"] = f"{num}. {desc} — NÃO ENCONTRADO"
    return mapping

# ── Pipeline ──────────────────────────────────────────────────────────────────

def _parse_question_numbers(text: str) -> set[int]:
    """Extract question numbers from a string like '1, 6, 10, 18, 31 e 74 (excluída a 25...)'."""
    if not text:
        return set()
    # Remove text in parentheses to avoid catching numbers inside notes
    cleaned = re.sub(r"\([^)]*\)", "", text)
    # Capture all digit sequences
    nums = re.findall(r"\b(\d{1,3})\b", cleaned)
    return set(int(n) for n in nums)


def _filter_relatorio_questoes(docx_path: Path, keep_numbers: set[int]) -> bool:
    """Edit a relatório docx to keep only the parecer of selected questions.

    Strategy: open docx with python-docx, identify section headers like
    'QUESTÃO N', 'Questão N', 'PARECER QUESTÃO N' etc., and remove paragraphs
    belonging to questions NOT in keep_numbers.
    """
    try:
        import docx as _docx
        doc = _docx.Document(str(docx_path))

        # Multiple regex patterns to catch question headers in various forms
        patterns = [
            re.compile(r"(?i)\bquest[ãa]o\s*[nNº°.:]*\s*(\d{1,3})\b"),
            re.compile(r"(?i)\bparecer\s+(?:da\s+)?(?:quest[ãa]o\s*)?[nNº°.:]*\s*(\d{1,3})\b"),
            re.compile(r"(?i)\b(?:relat[óo]rio|an[áa]lise)\s+(?:da\s+)?(?:quest[ãa]o\s*)?[nNº°.:]*\s*(\d{1,3})\b"),
            re.compile(r"(?i)^[Qq]\.?\s*(\d{1,3})\b"),
            re.compile(r"(?i)^(\d{1,3})[\.\)]\s+(?:[A-ZÁÉÍÓÚ])"),  # "10. CONTEÚDO..." style
        ]

        # Also detect strong/bold paragraphs as section markers
        def find_q_number(txt: str) -> int | None:
            txt = txt.strip()
            if not txt:
                return None
            for pat in patterns:
                m = pat.search(txt)
                if m:
                    try:
                        return int(m.group(1))
                    except (ValueError, IndexError):
                        continue
            return None

        current_q = None
        para_owner = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            # Check if this paragraph IS a question header
            qn = find_q_number(txt)
            # Heuristic: only treat as new section if the para is short (header-like)
            # OR it's bold/styled, OR the number found is at the very start
            if qn is not None:
                # If text length is short (<200) or starts with the question number
                if len(txt) < 200 or any(p.search(txt[:50]) for p in patterns):
                    current_q = qn
            para_owner.append(current_q)

        if not any(o in keep_numbers for o in para_owner if o is not None):
            log.warning("Filtro: nenhuma das questões a manter %s foi detectada no relatório",
                       sorted(keep_numbers))
            return False

        # Decide which paragraphs to remove
        to_remove_idx = []
        for i, owner in enumerate(para_owner):
            if owner is not None and owner not in keep_numbers:
                to_remove_idx.append(i)

        if not to_remove_idx:
            return False  # nothing to remove

        log.info("Filtro de relatório: removendo %d parágrafos (de %d total)",
                 len(to_remove_idx), len(doc.paragraphs))

        # Remove paragraphs in reverse to preserve indexes
        for i in reversed(to_remove_idx):
            p = doc.paragraphs[i]
            elem = p._element
            elem.getparent().remove(elem)

        doc.save(str(docx_path))
        return True
    except Exception as e:
        log.warning("Falha ao filtrar relatório %s: %s", docx_path.name, e)
        return False


def _convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convert DOCX to PDF using LibreOffice (soffice). Returns True on success.

    On Render free/Starter plans, LibreOffice is not installed by default.
    To enable PDF conversion, add to the Build Command:
        apt-get update && apt-get install -y libreoffice && pip install -r requirements.txt
    Or use Docker with a custom image.
    """
    import subprocess
    soffice_candidates = [
        "soffice", "libreoffice",
        "/usr/bin/soffice", "/usr/bin/libreoffice",
        "/opt/libreoffice/program/soffice",
        "/usr/lib/libreoffice/program/soffice",
    ]
    soffice_bin = None
    for cand in soffice_candidates:
        try:
            r = subprocess.run([cand, "--version"], capture_output=True, timeout=5)
            if r.returncode == 0:
                soffice_bin = cand
                log.info("LibreOffice encontrado: %s", cand)
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    if not soffice_bin:
        return False

    try:
        out_dir = pdf_path.parent
        result = subprocess.run(
            [soffice_bin, "--headless", "--convert-to", "pdf",
             "--outdir", str(out_dir), str(docx_path)],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode != 0:
            log.warning("soffice falhou: %s", result.stderr)
            return False
        generated = out_dir / (docx_path.stem + ".pdf")
        if generated.exists():
            if generated != pdf_path:
                generated.rename(pdf_path)
            return True
        return False
    except Exception as e:
        log.warning("Erro na conversão PDF: %s", e)
        return False


def process_zip(zip_path: Path, session_dir: Path, form_data: dict, api_key: str) -> dict:
    log.info("Pipeline start: %s", zip_path.name)

    # 1. Extract — handle Brazilian Portuguese filenames with cp437/latin-1 encoding
    extract_dir = session_dir / "extracted"
    extract_dir.mkdir()
    with zipfile.ZipFile(zip_path) as zf:
        for info in zf.infolist():
            # ZIP spec uses cp437 by default; many tools encode filenames as latin-1 or utf-8
            try:
                # Try to decode the raw bytes as utf-8 first, then fall back
                raw = info.filename.encode("cp437", errors="replace")
                try:
                    fixed_name = raw.decode("utf-8")
                except UnicodeDecodeError:
                    fixed_name = raw.decode("latin-1", errors="replace")
            except Exception:
                fixed_name = info.filename
            # Sanitize: ASCII-safe filename for filesystem
            safe_name = re.sub(r"[^\w\-./ ]", "_", fixed_name)
            target = extract_dir / safe_name
            if info.is_dir():
                target.mkdir(parents=True, exist_ok=True)
            else:
                target.parent.mkdir(parents=True, exist_ok=True)
                with zf.open(info) as src, open(target, "wb") as dst:
                    shutil.copyfileobj(src, dst)

    all_files = [f for f in extract_dir.rglob("*") if f.is_file()]
    file_list = [str(f.relative_to(extract_dir)) for f in all_files]
    log.info("Files: %s", file_list)

    # 2. Classify — apenas 4 categorias relevantes para a IA:
    #    (a) modelo da petição (.docx)
    #    (b) ficha do cliente (.xlsx)
    #    (c) procuração (.pdf) — para backup dos dados pessoais
    #    (d) relatórios técnicos das questões (.docx)
    # Todos os outros arquivos vão para "outros_arquivos" — entram no ZIP final
    # mas NÃO são lidos pela IA, economizando tokens e custos.

    model_keywords      = ("modelo", "peticao", "petição", "inicial")
    relatorio_keywords  = ("relatorio", "relatório", "questao", "questão",
                           "questoes", "questões", "parecer", "tecnico", "técnico",
                           "fundamentacao", "fundamentação")
    procuracao_keywords = ("procuracao", "procuração", "procuracoes", "procurações")
    ficha_keywords      = ("ficha",)

    docx_model = None
    docx_relatorios = []
    pdf_procuracao = None
    xlsx_ficha = None
    outros_arquivos = []  # vão pro ZIP mas não são lidos

    for fpath in all_files:
        nl = fpath.name.lower()
        if nl.endswith(".docx"):
            if any(k in nl for k in model_keywords):
                docx_model = fpath
            elif any(k in nl for k in relatorio_keywords):
                docx_relatorios.append(fpath)
            else:
                outros_arquivos.append(fpath)
        elif nl.endswith(".pdf"):
            if any(k in nl for k in procuracao_keywords):
                if pdf_procuracao is None:  # primeira procuração encontrada
                    pdf_procuracao = fpath
                else:
                    outros_arquivos.append(fpath)
            else:
                outros_arquivos.append(fpath)
        elif nl.endswith((".xlsx", ".xls")):
            if any(k in nl for k in ficha_keywords):
                if xlsx_ficha is None:
                    xlsx_ficha = fpath
                else:
                    outros_arquivos.append(fpath)
            else:
                outros_arquivos.append(fpath)
        else:
            outros_arquivos.append(fpath)

    log.info(
        "Classificação: modelo=%s, ficha=%s, procuração=%s, relatórios=%d, "
        "outros (não lidos)=%d",
        docx_model.name if docx_model else "AUSENTE",
        xlsx_ficha.name if xlsx_ficha else "AUSENTE",
        pdf_procuracao.name if pdf_procuracao else "AUSENTE",
        len(docx_relatorios),
        len(outros_arquivos),
    )

    if not docx_model:
        return {"error": "Nenhum arquivo .docx modelo encontrado no ZIP."}

    # 3. Build text block — apenas as 4 fontes essenciais
    parts = []

    parts.append("=== INVENTÁRIO DO ZIP ===\n" + "\n".join(f"  • {f}" for f in file_list))
    parts.append(f"""=== DADOS FORNECIDOS PELO USUÁRIO ===
Tipo de Ação: {form_data.get('tipo_acao','')}
Comarca: {form_data.get('comarca','')}
Resumo dos Fatos: {form_data.get('fatos','')}
Pedidos Adicionais: {form_data.get('pedidos','')}
Observações: {form_data.get('obs','')}""")

    # ─── FONTE 1/4: FICHA DO CLIENTE (XLSX) ────────────────────────────────────
    # Parsing estruturado + bloco de dados autoritativos
    ficha_estruturada = {}
    if xlsx_ficha:
        ficha_estruturada = parse_ficha_cliente(xlsx_ficha)

    if ficha_estruturada:
        fe = ficha_estruturada
        auth_block = ["\n=== DADOS AUTORITATIVOS DO CLIENTE (extraídos da Ficha — USE EXATAMENTE ESTES VALORES) ==="]

        # ─── DADOS DO CONCURSO (extraídos da ficha — fonte da verdade) ────────
        if fe.get("estado_concurso") or fe.get("instituicao_concurso") or fe.get("concurso"):
            auth_block.append("")
            auth_block.append("─── DADOS DO CONCURSO PRESTADO ───")
            if fe.get("instituicao_concurso"):
                auth_block.append(f"Instituição: {fe['instituicao_concurso']}")
            if fe.get("estado_concurso"):
                auth_block.append(
                    f"Estado: {fe['estado_concurso']} ({fe.get('uf_concurso','')})"
                )
            if fe.get("sigla_concurso"):
                auth_block.append(f"Sigla: {fe['sigla_concurso']}")
            if fe.get("concurso"):
                auth_block.append(f"Nome completo: {fe['concurso']}")
            auth_block.append("")

        if fe.get("nome_cliente"):
            auth_block.append(f"NOME COMPLETO: {fe['nome_cliente']}")
        if fe.get("rg"):
            auth_block.append(f"RG: {fe['rg']}")
        if fe.get("cpf"):
            auth_block.append(f"CPF: {fe['cpf']}")
        if fe.get("estado_civil"):
            auth_block.append(f"ESTADO CIVIL: {fe['estado_civil']}")
        if fe.get("profissao"):
            auth_block.append(f"PROFISSÃO: {fe['profissao']}")
        if fe.get("email"):
            auth_block.append(f"E-MAIL: {fe['email']}")
        if fe.get("endereco_completo"):
            auth_block.append(f"ENDEREÇO (residência do cliente): {fe['endereco_completo']}")
        if fe.get("cidade"):
            auth_block.append(f"CIDADE (residência): {fe['cidade']}")
        if fe.get("uf"):
            auth_block.append(f"UF (residência): {fe['uf']}")
        if fe.get("cep"):
            auth_block.append(f"CEP: {fe['cep']}")
        if fe.get("comarca"):
            auth_block.append(f"COMARCA (de endereçamento): {fe['comarca']}")
        if fe.get("concurso"):
            auth_block.append(f"CONCURSO (nome completo): {fe['concurso']}")
        if fe.get("sigla_concurso"):
            auth_block.append(f"SIGLA DO CONCURSO: {fe['sigla_concurso']}")
        if fe.get("instituicao_concurso"):
            auth_block.append(f"INSTITUIÇÃO DO CONCURSO: {fe['instituicao_concurso']}")
        if fe.get("estado_concurso"):
            auth_block.append(f"ESTADO DO CONCURSO: {fe['estado_concurso']} ({fe.get('uf_concurso','')})")
            auth_block.append(
                f"⚠️ ATENÇÃO: O concurso é de {fe.get('uf_concurso','')} "
                f"({fe['estado_concurso']}), NÃO infira outro estado a partir do endereço do cliente."
            )
        if fe.get("banca"):
            auth_block.append(f"BANCA: {fe['banca']}")
        if fe.get("cargo"):
            auth_block.append(f"CARGO: {fe['cargo']}")
        else:
            auth_block.append("CARGO: [NÃO INFORMADO — manter o do modelo]")
        if fe.get("tipo_prova"):
            auth_block.append(f"TIPO DE PROVA: Tipo {fe['tipo_prova'].replace('Tipo','').strip()}")
        if fe.get("pontuacao_obtida"):
            auth_block.append(f"PONTUAÇÃO OBTIDA: {fe['pontuacao_obtida']}")
        if fe.get("delta_anulacao"):
            auth_block.append(f"GANHO COM AS ANULAÇÕES (delta — NÃO É O TOTAL): {fe['delta_anulacao']}")
        if fe.get("pontuacao_final_calculada"):
            auth_block.append(
                f"⚠️ PONTUAÇÃO FINAL APÓS ANULAÇÕES (TOTAL): "
                f"{fe['pontuacao_final_calculada']}"
            )
        if fe.get("nota_corte"):
            auth_block.append(f"NOTA DE CORTE: {fe['nota_corte']}")
        else:
            auth_block.append("NOTA DE CORTE: [NÃO INFORMADA NA FICHA]")
        if fe.get("questoes_anular"):
            auth_block.append(f"QUESTÕES A ANULAR: {fe['questoes_anular']}")
        if fe.get("gratuidade"):
            auth_block.append(f"REQUER GRATUIDADE: {fe['gratuidade']}")
        if fe.get("eliminado"):
            auth_block.append(f"AUTOR ELIMINADO?: {fe['eliminado']}")
        if fe.get("tipo_acao"):
            auth_block.append(f"TIPO DE AÇÃO: {fe['tipo_acao']}")
        if fe.get("resumo_fatos"):
            auth_block.append(f"RESUMO DOS FATOS: {fe['resumo_fatos']}")
        auth_block.append(
            "\nREGRA: a petição final deve usar EXATAMENTE os valores acima — "
            "tanto para o cliente (nome, qualificação, endereço) quanto para o concurso "
            "(instituição, estado, sigla, banca, cargo, pontuação).\n"
            "Se um campo está [NÃO INFORMADO], mantenha o valor do modelo.\n"
        )
        parts.append("\n".join(auth_block))

    # Texto bruto da ficha (backup, caso o parser tenha perdido algo)
    if xlsx_ficha:
        parts.append(
            f"\n=== FICHA DO CLIENTE (texto bruto — backup) — arquivo: {xlsx_ficha.name} ===\n"
            + read_text_file(xlsx_ficha)
        )

    # ─── FONTE 2/4: PROCURAÇÃO (PDF) — backup de dados pessoais ──────────────
    if pdf_procuracao:
        proc_text = extract_pdf_text(pdf_procuracao, max_chars=10_000)
        parts.append(
            f"\n=== PROCURAÇÃO — arquivo: {pdf_procuracao.name} ===\n"
            f"Use APENAS para confirmar dados pessoais (nome, CPF, RG, endereço) "
            f"caso a ficha esteja incompleta.\n\n"
            f"{proc_text}"
        )

    # ─── FONTE 3/4: RELATÓRIOS TÉCNICOS DAS QUESTÕES ─────────────────────────
    for fpath in docx_relatorios:
        rname = fpath.relative_to(extract_dir)
        full_relat = read_docx_text(fpath, max_chars=200_000)
        parts.append(
            f"\n=== RELATÓRIO TÉCNICO DAS QUESTÕES — arquivo: {rname} ===\n"
            f"Este arquivo contém os pareceres técnicos de várias questões, mas você\n"
            f"deve usar APENAS as questões que estão na lista 'QUESTÕES A ANULAR' da ficha.\n"
            f"\n"
            f"INSTRUÇÕES PARA POPULAR O ARRAY 'questoes' DO JSON:\n"
            f"1. Para CADA questão da lista da ficha, localize o parecer correspondente\n"
            f"   neste relatório (procure por 'Questão N', 'PARECER QUESTÃO N', etc.).\n"
            f"2. Se o relatório tiver um tópico final 'RESUMOS' ou 'SÍNTESE' (geralmente\n"
            f"   no final do arquivo, com todos os enunciados + fundamentações compactos),\n"
            f"   USE PREFERENCIALMENTE essa versão para o campo 'relatorio_integra' —\n"
            f"   ela já vem otimizada para a petição inicial.\n"
            f"3. Caso contrário, use o parecer técnico completo da questão.\n"
            f"4. Para cada questão, popule:\n"
            f"   - numero, vicio (palavras-chave em CAIXA ALTA)\n"
            f"   - enunciado: comando + texto da questão num único parágrafo\n"
            f"   - alternativas: lista A/B/C/D/E\n"
            f"   - relatorio_integra: a fundamentação (do tópico Resumos, se existir)\n"
            f"5. Se uma questão da lista da ficha NÃO estiver no relatório, adicione-a\n"
            f"   ao array 'dados_ausentes' como 'Relatório técnico da questão N ausente'.\n"
            f"\n"
            f"NÃO gere pares de 'substituicoes' para os blocos das questões — o sistema\n"
            f"reescreve esse capítulo automaticamente usando o array 'questoes'.\n\n"
            f"{full_relat}"
        )

    # ─── FONTE 4/4: MODELO DA PETIÇÃO ────────────────────────────────────────
    modelo_text = read_docx_text(docx_model, max_chars=200_000)
    parts.append(
        f"\n=== MODELO DA PETIÇÃO (texto integral) — arquivo: {docx_model.name} ===\n"
        f"Identifique TODOS os trechos que se referem ao CLIENTE ANTIGO\n"
        f"(nome, RG, CPF, endereço, pontuação, banca, cargo, comarca, etc.)\n"
        f"e gere pares de 'buscar' (texto antigo) → 'substituir' (valor correto da ficha).\n"
        f"\n"
        f"⚠️ NÃO gere pares de substituição para os BLOCOS DE QUESTÕES no capítulo\n"
        f"'DO ROL DE QUESTÕES ILEGAIS' (ou equivalente). O sistema vai reescrever\n"
        f"esse capítulo automaticamente usando os dados do array 'questoes' do JSON.\n"
        f"\n"
        f"Sua tarefa para as questões é APENAS popular o array 'questoes' com:\n"
        f"  - numero: número da questão (conforme a ficha)\n"
        f"  - vicio: tipo do vício (ERRO GROSSEIRO, EXTRAPOLAÇÃO DO EDITAL, etc.)\n"
        f"  - enunciado: comando + enunciado da questão num único parágrafo\n"
        f"  - alternativas: lista das alternativas (A, B, C, D, E)\n"
        f"  - relatorio_integra: a fundamentação COMPLETA da questão extraída\n"
        f"    do RELATÓRIO TÉCNICO (não invente — use o texto real do relatório).\n"
        f"\n"
        f"⚠️ APENAS no rol de pedidos finais (frase tipo 'declarar a nulidade das\n"
        f"questões impugnadas 10, 25, 31 e 34'), você deve sim gerar par de\n"
        f"substituição com os números novos da ficha do cliente.\n"
        f"\n"
        f"⚠️ REGRA ANTI-RECURSÃO: NUNCA gere um par onde 'buscar' está contido\n"
        f"dentro de 'substituir'. Exemplo PROIBIDO:\n"
        f"  buscar: 'PCES'\n"
        f"  substituir: 'PCES (informação adicional)'\n"
        f"O sistema rejeita pares assim por causarem loop infinito de substituições.\n\n"
        f"{modelo_text}"
    )

    full_text = "\n\n".join(parts)
    log.info("Text block: %d chars (budget: %d) — fontes: ficha=%s, procuração=%s, "
             "relatórios=%d, modelo=%s, ignorados=%d",
             len(full_text), TOTAL_CHAR_BUDGET,
             "sim" if xlsx_ficha else "não",
             "sim" if pdf_procuracao else "não",
             len(docx_relatorios),
             "sim" if docx_model else "não",
             len(outros_arquivos))

    # Free memory before API call
    del parts
    gc.collect()

    # 4. Call Claude (text only)
    data = call_claude(api_key, full_text)
    log.info("Claude OK")
    del full_text
    gc.collect()

    # 5. Unpack modelo
    unpacked_dir = session_dir / "unpacked"
    if not unpack_docx(docx_model, unpacked_dir):
        return {"error": "Falha ao desempacotar o modelo DOCX."}

    # 6. Apply edits
    changes = apply_substitutions(unpacked_dir, data)

    # 6b. ROUND 2 REVIEW — extract current text, ask Claude what data is still old
    try:
        doc_xml_path = unpacked_dir / "word" / "document.xml"
        if doc_xml_path.exists():
            current_xml = doc_xml_path.read_text(encoding="utf-8")
            current_text = re.sub(r"<[^>]+>", " ", current_xml)
            current_text = re.sub(r"\s+", " ", current_text).strip()
            current_text_limited = current_text[:200_000]

            # Build authoritative client data block — prefer ficha values
            auth_lines = ["DADOS CORRETOS DO CLIENTE (use EXATAMENTE estes valores):"]
            if ficha_estruturada:
                fe = ficha_estruturada
                if fe.get("nome_cliente"):
                    auth_lines.append(f"- Nome: {fe['nome_cliente']}")
                if fe.get("comarca"):
                    auth_lines.append(f"- Comarca: {fe['comarca']}")
                if fe.get("banca"):
                    auth_lines.append(f"- Banca: {fe['banca']}")
                if fe.get("cargo"):
                    auth_lines.append(f"- Cargo: {fe['cargo']}")
                if fe.get("tipo_prova"):
                    auth_lines.append(f"- Tipo de prova: {fe['tipo_prova']}")
                if fe.get("pontuacao_obtida"):
                    auth_lines.append(f"- Pontuação obtida: {fe['pontuacao_obtida']}")
                if fe.get("pontuacao_final_calculada"):
                    auth_lines.append(f"- Pontuação final após anulações: {fe['pontuacao_final_calculada']}")
                if fe.get("nota_corte"):
                    auth_lines.append(f"- Nota de corte: {fe['nota_corte']}")
                if fe.get("questoes_anular"):
                    auth_lines.append(f"- Questões a anular: {fe['questoes_anular']}")
            cliente_novo = data.get("cliente", {})
            if cliente_novo.get("rg"):
                auth_lines.append(f"- RG: {cliente_novo['rg']}")
            if cliente_novo.get("cpf"):
                auth_lines.append(f"- CPF: {cliente_novo['cpf']}")
            if cliente_novo.get("endereco"):
                auth_lines.append(f"- Endereço: {cliente_novo['endereco']}, {cliente_novo.get('cidade','')}/{cliente_novo.get('uf','')}")

            review_prompt = f"""{chr(10).join(auth_lines)}

TAREFA: Analise o TEXTO DA PETIÇÃO abaixo. Identifique APENAS trechos que CONTRADIGAM
diretamente os DADOS CORRETOS acima (ou seja, dados de outro cliente que ainda restaram).

REGRAS RIGOROSAS:
- Só gere par de substituição se o texto da petição contém um valor DIFERENTE do correto.
- NÃO gere pares se o texto já está correto (mesmo que pareça simplificado).
- NÃO gere pares para o cabeçalho de endereçamento já correto.
- NÃO substitua valores corretos por valores semelhantes.
- Se o texto está correto: retorne {{"substituicoes": []}}.
- Cada par precisa ter "buscar" e "substituir" como strings NÃO VAZIAS.

Para cada inconsistência ENCONTRADA, gere {{"buscar": "texto literal exato", "substituir": "valor correto"}}.

Retorne APENAS JSON puro — sem markdown, sem ```, sem explicações.

TEXTO DA PETIÇÃO ATUAL:
{current_text_limited}"""

            log.info("Round 2: enviando %d chars para revisão", len(review_prompt))
            client = anthropic.Anthropic(api_key=api_key, timeout=300.0, max_retries=2)
            review_msg = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=32000,
                system="Você é um revisor jurídico extremamente atento. Sua única tarefa é identificar inconsistências de dados entre uma petição editada e os dados corretos do cliente. Retorne APENAS JSON puro.",
                messages=[{"role": "user", "content": review_prompt}]
            )
            review_raw = "".join(b.text for b in review_msg.content if hasattr(b, "text"))
            log.info("Round 2 response: %d output tokens, raw preview: %s",
                     review_msg.usage.output_tokens if review_msg.usage else 0,
                     review_raw[:300])

            try:
                review_data = _parse_claude_json(review_raw)
                review_pairs_raw = review_data.get("substituicoes", [])
                # Filter out empty pairs
                review_pairs = [
                    p for p in review_pairs_raw
                    if isinstance(p, dict)
                    and p.get("buscar","").strip()
                    and p.get("substituir","").strip()
                    and p.get("buscar","").strip() != p.get("substituir","").strip()
                ]
                if review_pairs:
                    log.info("Round 2: %d pares válidos adicionais para aplicar (de %d retornados)",
                             len(review_pairs), len(review_pairs_raw))
                    review_changes = apply_substitutions(
                        unpacked_dir,
                        {"substituicoes": review_pairs, "processo": {"gratuidade": True}}
                    )
                    changes.extend(["📝 ROUND 2 (revisão automática):"] + review_changes)
                else:
                    log.info("Round 2: nenhum par válido (recebidos %d, todos vazios/idênticos)",
                             len(review_pairs_raw))
                    changes.append("✅ ROUND 2: nenhum dado antigo restante detectado")
            except Exception as e:
                log.warning("Round 2 review parse failed: %s. Raw: %s", e, review_raw[:500])
                changes.append(f"⚠️ Round 2 falhou ao parsear: {e}")
    except Exception as e:
        log.warning("Round 2 skipped: %s", e)
        changes.append(f"⚠️ Round 2 não executou: {e}")

    # 7. Repack
    safe_nome  = re.sub(r"[^\w\s]","",data.get("cliente",{}).get("nome_completo","Cliente")).replace(" ","_")
    tipo_acao  = re.sub(r"[^\w\s]","",data.get("processo",{}).get("tipo_acao","Peticao")).replace(" ","_")
    data_hoje  = datetime.now().strftime("%Y-%m-%d")
    out_name   = f"Peticao_{tipo_acao}_{safe_nome}_{data_hoje}.docx"
    out_path   = session_dir / out_name

    if not pack_docx(unpacked_dir, out_path, docx_model):
        return {"error": "Falha ao reempacotar o DOCX final."}

    # 8. Delivery dir
    delivery = session_dir / "entrega"
    delivery.mkdir()
    shutil.copy(out_path, delivery / out_name)
    for fpath in all_files:
        if fpath != docx_model:
            dest = delivery / fpath.name
            if not dest.exists():
                shutil.copy(fpath, dest)

    # 8b. Filter relatórios técnicos: keep only questions in 'questoes_anular'
    if ficha_estruturada and ficha_estruturada.get("questoes_anular"):
        questoes_a_manter = _parse_question_numbers(ficha_estruturada["questoes_anular"])
        log.info("Questões a manter no relatório: %s", questoes_a_manter)
        for fpath in docx_relatorios:
            delivery_path = delivery / fpath.name
            if delivery_path.exists():
                try:
                    filtered = _filter_relatorio_questoes(delivery_path, questoes_a_manter)
                    if filtered:
                        changes.append(f"📝 Relatório '{fpath.name}' filtrado: mantidas questões {sorted(questoes_a_manter)}")
                except Exception as e:
                    log.warning("Falha ao filtrar relatório %s: %s", fpath.name, e)
                    changes.append(f"⚠️ Não foi possível filtrar '{fpath.name}': {e}")

    # 9. Rename
    rol = data.get("rol_documentos", [])
    rename_map = rename_files_by_rol(delivery, rol) if rol else {}

    # 9b. Convert petição DOCX to PDF
    pdf_name = out_name.replace(".docx", ".pdf")
    pdf_path = delivery / pdf_name
    # Locate the renamed petição docx
    docx_in_delivery = None
    for f in delivery.iterdir():
        if f.suffix.lower() == ".docx" and (f.name == out_name or "peticao" in f.name.lower()):
            docx_in_delivery = f
            break
    if docx_in_delivery:
        try:
            if _convert_to_pdf(docx_in_delivery, pdf_path):
                changes.append(f"📄 PDF gerado: {pdf_name}")
            else:
                changes.append(f"⚠️ Não foi possível gerar PDF (LibreOffice indisponível)")
        except Exception as e:
            log.warning("Falha na conversão PDF: %s", e)
            changes.append(f"⚠️ Falha ao gerar PDF: {e}")

    # 10. Final zip
    zip_name = f"Entrega_{safe_nome}_{data_hoje}.zip"
    zip_out  = OUTPUT_DIR / zip_name
    with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zf:
        for fpath in delivery.iterdir():
            if fpath.is_file():
                zf.write(fpath, fpath.name)

    log.info("Done: %s", zip_name)
    return {
        "success":        True,
        "zip_filename":   zip_name,
        "docx_filename":  out_name,
        "cliente":        data.get("cliente", {}),
        "processo":       data.get("processo", {}),
        "questoes":       [{"numero": q.get("numero"), "vicio": q.get("vicio")} for q in data.get("questoes", [])],
        "changes":        changes,
        "rename_map":     rename_map,
        "dados_ausentes": data.get("dados_ausentes", []),
        "relatorio":      data.get("relatorio_alteracoes", []),
    }

# ── Global error handlers ─────────────────────────────────────────────────────

@app.errorhandler(Exception)
def handle_exception(e):
    log.exception("Unhandled exception")
    return jsonify({"error": f"Erro interno: {str(e)}"}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "Arquivo muito grande. Limite: 100 MB."}), 413

# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/health")
def health():
    """Healthcheck + diagnostic endpoint."""
    import sys
    checks = {}
    try:
                checks["pypdf"] = "ok"
    except ImportError as e:
        checks["pypdf"] = f"MISSING: {e}"
    try:
        import openpyxl
        checks["openpyxl"] = "ok"
    except ImportError as e:
        checks["openpyxl"] = f"MISSING: {e}"
    try:
        import anthropic
        checks["anthropic"] = "ok"
    except ImportError as e:
        checks["anthropic"] = f"MISSING: {e}"
    checks["unpack_py"]  = "ok" if (SCRIPTS / "unpack.py").exists() else "MISSING"
    checks["pack_py"]    = "ok" if (SCRIPTS / "pack.py").exists() else "MISSING"
    checks["upload_dir"] = str(UPLOAD_DIR)
    checks["output_dir"] = str(OUTPUT_DIR)
    checks["python"]     = sys.version
    return jsonify(checks)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/gerar", methods=["POST"])
def gerar():
    import traceback
    try:
        api_key = request.form.get("api_key", "").strip()
        if not api_key:
            return jsonify({"error": "Chave de API não fornecida."}), 400
        if "zip_file" not in request.files or request.files["zip_file"].filename == "":
            return jsonify({"error": "Arquivo ZIP não enviado."}), 400
        zip_file = request.files["zip_file"]
        if not zip_file.filename.lower().endswith(".zip"):
            return jsonify({"error": "Apenas arquivos .zip são aceitos."}), 400

        session_id  = str(uuid.uuid4())
        session_dir = UPLOAD_DIR / session_id
        session_dir.mkdir(parents=True, exist_ok=True)
        zip_path = session_dir / "input.zip"

        log.info("Saving uploaded ZIP...")
        zip_file.save(str(zip_path))
        log.info("ZIP saved: %s bytes", zip_path.stat().st_size)

        form_data = {k: request.form.get(k, "") for k in ("tipo_acao","comarca","fatos","pedidos","obs")}

        try:
            result = process_zip(zip_path, session_dir, form_data, api_key)
        except anthropic.AuthenticationError:
            return jsonify({"error": "Chave de API inválida. Verifique suas credenciais em console.anthropic.com"}), 401
        except anthropic.RateLimitError:
            return jsonify({"error": "Limite de uso da API Anthropic atingido. Aguarde alguns minutos."}), 429
        except anthropic.APIStatusError as e:
            log.exception("Anthropic API error")
            return jsonify({"error": f"Erro da API Anthropic: {e.status_code} — {e.message}"}), 502
        except json.JSONDecodeError as e:
            log.exception("JSON decode error")
            return jsonify({"error": f"Claude não retornou JSON válido: {e}"}), 500
        except Exception as e:
            log.exception("Pipeline error")
            tb = traceback.format_exc()
            return jsonify({"error": str(e), "traceback": tb[-2000:]}), 500

        if "error" in result:
            return jsonify(result), 500

        result["session_id"] = session_id
        log.info("Request completed OK: %s", result.get("zip_filename"))
        return jsonify(result)

    except Exception as e:
        log.exception("Outer exception in /gerar")
        tb = traceback.format_exc()
        return jsonify({"error": f"Erro inesperado: {str(e)}", "traceback": tb[-2000:]}), 500

@app.route("/download/<session_id>/<filename>")
def download(session_id, filename):
    if not re.match(r"^[\w\-]+$", session_id):
        return "ID inválido", 400
    for base in [OUTPUT_DIR, UPLOAD_DIR / session_id]:
        path = base / filename
        if path.exists():
            return send_file(str(path), as_attachment=True, download_name=filename)
    return "Arquivo não encontrado", 404

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
